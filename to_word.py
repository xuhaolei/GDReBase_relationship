# -*- encoding:utf-8 -*-
import csv

from random import randint
from docx import Document
from docx.shared import RGBColor

from match import DiseaseMatch,BacteriumMatch
from options import args

class Atten2Docx(object):
    """docstring for Atten2Docx"""
    def __init__(self):
        super(Atten2Docx, self).__init__()
        self.doc = Document()
    def add_text(self,text,label):
        assert len(text)==len(label),"text和attn长度不一致"
        p = self.doc.add_paragraph()
        for i in range(len(text)):
            word = p.add_run(text[i])
            if label[i] != 'O':
                color = word.font.color.rgb = RGBColor(255, 0, 0)
                word.font.underline = True
            # else:
            #     word.font.color.rgb = RGBColor(*color)

    def save_file(self,file_name):
        if file_name[-5:]!='.docx':
            file_name = file_name + '.docx'
        self.doc.save(file_name)

# if args.dataset == 'disease':
#     Match = DiseaseMatch
# elif args.dataset == 'bacerium':
#     Match = BacteriumMatch
# else:
#     raise Exception

# match = Match()
# match.build_tree()

# doc = Atten2Docx(0.5)

# with open('data/test_disease.csv','r',encoding='utf-8') as f:
#     lines = csv.reader(f)
#     text = ''
#     labels = ''
#     for cnt,line in enumerate(lines):
#         labels,tokens,entities = match.paper_entities(line[2].strip().lower())
        
#         labels_set = set()
#         for label in labels:
#             for idx in range(label[0],label[1]+1):
#                 labels_set.add(idx)
#         paper = ' '.join(tokens)

#         labels_one_hot = ''
#         first = True
#         for idx,token in enumerate(tokens):
#             if not first: labels_one_hot += 'O'
#             if idx in labels_set:
#                 labels_one_hot += 'X'*len(token)
#             else:
#                 labels_one_hot += 'O'*len(token)
#             first = False
#         doc.add_text(paper,labels_one_hot)
#         if cnt >= 9:
#             break

# doc.save_file('paper')
